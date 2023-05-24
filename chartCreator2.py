# from pydoc import doc
import sys
import os
import docx
from docx import Document
from docx.shared import Inches

import datetime

from PyQt5.QtCore import *
from PyQt5 import QtCore
from PyQt5.QtGui import *
from PyQt5 import QtGui
from PyQt5.QtGui import QPixmap
from PyQt5.QtWidgets import *
from PyQt5 import QtWidgets
from PyQt5.QtCore import QSize
from PyQt5.QtGui import QIcon
from numpy import save


class ClassUi(object):
    def setup(self, MainW):
        MainW.setObjectName("MainW")
        MainW.resize(800, 745)
        self.setWindowTitle("Chart Creator")
 

        self.mainlayout = QtWidgets.QVBoxLayout()

        self.centralwidget = QtWidgets.QWidget(MainW)
        # sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
        # sizePolicy.setHorizontalStretch(0)
        # sizePolicy.setVerticalStretch(0)
        # sizePolicy.setHeightForWidth(self.centralwidget.sizePolicy().hasHeightForWidth())
        # self.centralwidget.setSizePolicy(sizePolicy)
        self.centralwidget.setLayout(self.mainlayout)
        MainW.setCentralWidget(self.centralwidget)

    # direcwidget is a container widget for the top part of the program where a directory is chosen
        self.direcwidget = QtWidgets.QWidget()
        self.direclayout = QtWidgets.QGridLayout()
        # self.direcwidget.setLayout(self.direclayout)
        self.mainlayout.addWidget(self.direcwidget)

        # Button that opens directory dialog when pressed
        self.opdirbut = QtWidgets.QPushButton()
        self.opdirbut.setText("Choose")
        self.opdirbut.setFixedSize(65, 34)

        # Line Edit that displays current directory
        self.linepath = QtWidgets.QLineEdit()
        self.linepath.setText(os.getcwd())

        # Button that changes directory to parent directory of the current one
        self.backpath = QtWidgets.QPushButton()
        self.backpath.setFixedSize(20, 20)
        self.backpath.setText("^")

        self.createChart = QtWidgets.QPushButton()
        self.createChart.setFixedSize(100, 34)
        self.createChart.setText("Create")
        self.createChart.setEnabled(False)

        self.widthLabel = QtWidgets.QLabel(self)
        self.widthLabel.setText("Width (Inch) :")
        self.widthLine = QLineEdit(self)
        self.widthLine.setValidator(QDoubleValidator(0.99, 99.9, 9))

        self.heightLabel = QtWidgets.QLabel(self)
        self.heightLabel.setText("Height (Inch) :")
        self.heightLine = QLineEdit(self)
        self.heightLine.resize(100, 32)
        self.heightLine.setValidator(QDoubleValidator(0.99, 99.9, 9))

    # Positioning of widgets inside widget container direcwidget
        # self.direclayout.addWidget(self.opdirbut, 0, 0, 2, 1)
        # self.direclayout.addWidget(self.linepath, 0, 2, 1, 3)
        # self.direclayout.addWidget(self.backpath, 1,4, 1, 1)

    # Scrollwidget is the area wherein a container widget widgetforscroll holds all image icons in a grid
        self.scrollwidget = QtWidgets.QScrollArea()
        self.mainlayout.addWidget(self.scrollwidget)
        self.scrollwidget.setWidgetResizable(True)
        self.scrollwidget.setVerticalScrollBarPolicy(
            QtCore.Qt.ScrollBarAlwaysOn)

        self.scrollgrid = QtWidgets.QGridLayout()

        self.widgetforscroll = QtWidgets.QWidget()
        self.widgetforscroll.setLayout(self.scrollgrid)

        self.scrollwidget.setWidget(self.widgetforscroll)
        # self.mainlayout.addWidget(self.widthLabel)
        # self.mainlayout.addWidget(self.widthLine)
        # self.mainlayout.addWidget(self.heightLabel)
        # self.mainlayout.addWidget(self.heightLine)
        self.mainlayout.addWidget(self.createChart)


# Contains logic of program
class MainWindow(QtWidgets.QMainWindow, ClassUi):
    def __init__(self):
        super().__init__()
        self.setup(self)
        self.createActions()
        self.createMenuBar()
        self.connectActions()
        self.setAcceptDrops(True)
        self.counter = 0
        global files
        files = []

    # Counter variables for keeping track of where to layout items
        self.picturerow = 0
        self.picturecolumn = 0
        self.howmany = 0

    # Assigns class methods to directory buttons
        self.opdirbut.clicked.connect(self.opdial)
        # self.backpath.clicked.connect(self.upath)
        self.createChart.clicked.connect(self.createChartAction)

    def createMenuBar(self):
        menuBar = QMenuBar(self)
        fileMenu = QMenu("&File", self)
        menuBar.addMenu(fileMenu)
        fileMenu.addAction(self.newAction)
        fileMenu.addSeparator()
        fileMenu.addAction(self.exitAction)
        self.setMenuBar(menuBar)

    def createActions(self):
        self.newAction = QAction(self)
        self.newAction.setText("&New")
        self.exitAction = QAction("&Exit", self)

    def newFile(self):
        # Add logic to delete the old pics or empty the image views
        print("New file will open")
        files.clear()
        self.createChart.setEnabled(False)
        for i in reversed(range(self.scrollgrid.count())):
                widgetToRemove = self.scrollgrid.itemAt(i).widget()
                # remove it from the layout list
                self.scrollgrid.removeWidget(widgetToRemove)
                # remove it from the gui
                widgetToRemove.setParent(None)


    
    def exitWindow(self):
        # Add logic to exit the application
        print("Exiting the window")
        quit()

    def connectActions(self):
        self.newAction.triggered.connect(self.newFile)
        self.exitAction.triggered.connect(self.exitWindow)


# Each time this function is called, a new widget called newwidget is created containing "pic" in a pixmap label and a text label and layed out
# on the widgetforscroll widget through its scrollgrid layout. Each time the function is called, picture column increments by one at the end of the function
# when all the columns in a row are filled, picture column is reset to 0 and and picture row is incremented. Picture row and picture column are used in positioning
# the newwidgets in the scrollgrid.

    def addpicture(self, pic):
        if self.picturecolumn == 3:
            self.picturecolumn = 0
            self.picturerow += 1
        self.howmany += 1

# newwidget is object of picwidg class containing pixmap and text label
        newwidget = picwidg(self.howmany, pic)

# This function was not required to be created, it was only created for the purpose of the Qtimer singleshot implementation.
# The newwidget is being positioned on the scrollgrid layout here.
        def addnewone(lyout, nw, rw, cl):
            lyout.addWidget(nw, rw, cl)

        QtCore.QTimer.singleShot(
            self.howmany*10,
            lambda sc=self.scrollgrid, 
            nr=newwidget, ow=self.picturerow, 
            mn=self.picturecolumn: addnewone(
                sc, nr, ow, mn)
        )
# Incrementing column by 1 for the next time function is called
        self.picturecolumn += 1

# This is the function connected to the choose dialog button. It opens a QFileDialog window which allows you to only choose a directory folder.
# When the folder is chosen:
        # 1: The linepath text is set the to the new directory
        # 2: Any previous picwidg objects are cleared from the scrollgrid layout
        # 3: Picture column and picture row variables are reset for positioning
        # 4: A for loop scans the new directory for files with .jpg or .png extensions
        # 5: The addpicture method is called with the filename as the argument
    def opdial(self):
        dialogbox = dialog()

        try:
            os.chdir(dialogbox.getExistingDirectory(
                options=QtWidgets.QFileDialog.DontUseNativeDialog))
            self.linepath.setText(os.getcwd())

            for i in reversed(range(self.scrollgrid.count())):
                widgetToRemove = self.scrollgrid.itemAt(i).widget()
                # remove it from the layout list
                self.scrollgrid.removeWidget(widgetToRemove)
                # remove it from the gui
                widgetToRemove.setParent(None)

            self.picturecolumn = 0
            self.picturerow = 0
            self.howmany = 0

            for a, b, c in os.walk(os.getcwd()):
                for i in c:
                    print(i)
                    print("Printing c", c)
                    if i[-4:].lower() == ".png" or i[-4:].lower() == ".jpg":
                        self.addpicture(i)

        except:
            pass

    def createChartAction(self):
        print(files)
        flat_files = sum(files, [])
        print("Flat Files", flat_files)
        folder_name = os.path.dirname(flat_files[0])
        print(folder_name)
        document = Document()
        section = document.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "\t\tNava Bharat Xerox"
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        for f in flat_files:
            run.add_picture(f, width=Inches(3.5), height=Inches(3.0))
            run.add_text("       ")
        now = datetime.datetime.now()
        filename = now.strftime("%d-%m%I-%M%p")
        save_filename = folder_name + '/' + str(filename)
        document.save(save_filename+".docx")
        print("File saved to "+save_filename)
        print("Document Created")
        os.startfile(save_filename+".docx")


    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        if self.counter == 0:
            first_files = [u.toLocalFile() for u in event.mimeData().urls()]
            print("Printing files first time", first_files)
            files.append(first_files)
            self.counter += 1
            for f in first_files:
                # print(f)
                # if i[-4:].lower() == ".png" or i[-4:].lower() == ".jpg":
                self.addpicture(f)
                self.createChart.setEnabled(True)
        else:
            second_files = [u.toLocalFile() for u in event.mimeData().urls()]
            print("Printing files second time", second_files)
            files.append(second_files)
            for f in second_files:
                # print(f)
                # if i[-4:].lower() == ".png" or i[-4:].lower() == ".jpg":
                self.addpicture(f)
                self.createChart.setEnabled(True)


# This is the class where newwidget instances are created
# Here 2 labels are created, one for the image, one for the text and packed in a vertical layout
class picwidg(QtWidgets.QWidget):
    whoshover = None
    picwidglist = []

    def __init__(self, numb, pic):
        super().__init__()
        self.setMouseTracking(True)
        self.numb = numb
        self.pic = pic
        picwidg.picwidglist.append(self)

        # SizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)

        newwidgetlayout = QtWidgets.QVBoxLayout()
        self.setLayout(newwidgetlayout)
        # self.setSizePolicy(SizePolicy)
        self.setMinimumSize(QtCore.QSize(200, 200))
        self.setMaximumSize(QtCore.QSize(450, 450))

        # Pic Label
        self.newpic = QtWidgets.QLabel()
        QtCore.QTimer.singleShot(self.numb*200, self.addingnewpic)
        self.newpic.setScaledContents(True)
        # self.newpic.setSizePolicy(SizePolicy)
        self.newpic.setGeometry(0, 0, 500, 500)
        # self.newpic.setStyleSheet("border:1px solid gray")

        # Picture text label
        self.newtext = QtWidgets.QLabel()
        font_metrics = QtGui.QFontMetrics(self.font())
        self.newtext.setAlignment(QtCore.Qt.AlignCenter)
        elided_text = font_metrics.elidedText(pic, QtCore.Qt.ElideRight, 200)
        self.newtext.setText(elided_text)

        newwidgetlayout.addWidget(self.newpic)
        newwidgetlayout.addWidget(self.newtext)

    def addingnewpic(self):
        self.newpic.setPixmap(QtGui.QPixmap(self.pic))


# Class for QFileDialog for selecting only directories
class dialog(QtWidgets.QFileDialog):
    def __init__(self):
        super().__init__()
        self.setFileMode(QtWidgets.QFileDialog.DirectoryOnly)
        self.setOption(QtWidgets.QFileDialog.DontUseNativeDialog, True)
        self.setOption(QtWidgets.QFileDialog.ShowDirsOnly, False)


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    w = MainWindow()
    w.show()
    sys.exit(app.exec_())
